import os
# from datetime import date
# from datetime import timedelta
from datetime import *

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt
from matplotlib import pyplot as plt
from scipy.optimize import curve_fit

from TimeStop import timeBreak
from TimeWork import timeWork
from change_pressure_interval import change_pressure_interval
from check_pressure import check_pressure
from concat_utc import concat_utc


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    # выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    # обозначаем конец позиции вывода
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


# функция для фита графика скорости счета к разности давлений
def linear(x, a, b):
    return a * x + b


# Длиннокод того, как строяться графики и заполняется ворд
def secProccesing(stday, stmonth, styear, endday, endmonth, endyear, path, pathpic, file1cl, file2cl,
                  proccessing_pressure, file_mask):
    datedirect = pathpic + '/{}'.format(styear)
    if ~os.path.exists(datedirect):
        try:
            os.mkdir(datedirect)
        except OSError:
            print("Создать директорию %s не удалось" % datedirect)
        else:
            print("Успешно создана директория %s " % datedirect)
    your_format = lambda x: '{{:0.{}f}}'.format(2 if x > 1 else 3).format(x)
    pd.set_option('display.float_format', your_format)
    # a = date(styear, stmonth, stday)
    # b = date(endyear, endmonth, endday)

    worktime = timeWork(stday, endday, styear, endyear, stmonth, endmonth, file1cl)

    timestop = timeBreak(stday, endday, styear, endyear, stmonth, endmonth, file1cl)
    # Сделать вывод остановок

    failstr_begin = []
    failstr_end = []
    lost_minutes = []
    for i in range(len(timestop.index)):
        failstr_begin.append(" {}  {}".format(timestop['stDate'][i].date(), timestop['stTime'][i]))
        failstr_end.append(" {}  {}".format(timestop['endDate'][i].date(), timestop['endTime'][i]))
        lost_minutes.append((timestop['endDate'][i].day - timestop['stDate'][i].day) * 24 * 60 + int(
            timestop['endTime'][0].split(':')[0]) * 60 + int(
            timestop['endTime'][0].split(':')[1]) - int(
            timestop['stTime'][0].split(':')[1]) - int(
            timestop['stTime'][0].split(':')[0]) * 60)

    font = {'weight': 'bold',
            'size': 14}

    plt.rc('font', **font)

    merge_utc = concat_utc(stday, endday, styear, endyear, stmonth, endmonth, file1cl)
    dont_use_it, cutting_intervals_uragan = check_pressure(stday, endday, styear, endyear, stmonth, endmonth,
                                                           file2cl)
    cutting_len = []
    new_time = []
    for g in range(len(merge_utc['time'])):
        new_time.append(datetime(int(str(merge_utc['DATE'][g]).split(' ')[0].split('-')[0]),
                                 int(str(merge_utc['DATE'][g]).split(' ')[0].split('-')[1]),
                                 int(str(merge_utc['DATE'][g]).split(' ')[0].split('-')[2]),
                                 int(merge_utc['time'][g].split(':')[0]),
                                 int(merge_utc['time'][g].split(':')[1])))
    merge_utc['new_format_time'] = new_time
    if len(cutting_intervals_uragan) != 0:
        for i in range(1, 5):
            for a in cutting_intervals_uragan:
                merge_utc['Nn%s' % i] = merge_utc['Nn%s' % i].where(
                    (merge_utc['new_format_time'] < a[0]) | (
                            merge_utc['new_format_time'] > a[1]), 0)
                merge_utc['N_noise%s' % i] = merge_utc['N_noise%s' % i].where(
                    (merge_utc['new_format_time'] < a[0]) | (
                            merge_utc['new_format_time'] > a[1]), 0)

        for a in cutting_intervals_uragan:
            cutting_len.append(len(merge_utc.index) - len(merge_utc[(merge_utc['new_format_time'] < a[0]) | (
                    merge_utc['new_format_time'] > a[1])].index))

    mean_pressure_uragan = change_pressure_interval(stday, endday, styear, endyear, stmonth, endmonth,
                                                    file2cl, cutting_len)
    mask_checker = False
    try:
        for i in range(1, 5):
            mask = pd.read_csv(
                '{}\\Maska_detAll{}.dat'.format(file_mask, i),
                sep='\t', names=['start', 'end'])
            mask_without_periods = mask.fillna(0)[mask.fillna(0)['end'] == 0][['start']].reset_index(drop=True)
            period_mask = mask.dropna().reset_index(drop=True)

            # выделение одиночных выбросов
            start_time = []
            start_date = []
            for j in mask_without_periods['start']:
                start_time.append(j.split(' ')[1])
                start_date.append(j.split(' ')[0].split('.')[2] + '-' + j.split(' ')[0].split('.')[1] + '-' +
                                  j.split(' ')[0].split('.')[0])
            mask_without_periods['start_time'] = start_time
            mask_without_periods['start_date'] = start_date
            mask_without_periods['start_date'] = pd.to_datetime(mask_without_periods['start_date'])

            # выделение периодов
            start_date = []
            end_date = []
            for k in range(len(period_mask['start'])):
                end_date.append(datetime(int(period_mask['end'][k].split(' ')[0].split('.')[2]),
                                         int(period_mask['end'][k].split(' ')[0].split('.')[1]),
                                         int(period_mask['end'][k].split(' ')[0].split('.')[0]),
                                         int(period_mask['end'][k].split(' ')[1].split(':')[0]),
                                         int(period_mask['end'][k].split(' ')[1].split(':')[1]),
                                         int(period_mask['end'][k].split(' ')[1].split(':')[2])))
                start_date.append(datetime(int(period_mask['start'][k].split(' ')[0].split('.')[2]),
                                           int(period_mask['start'][k].split(' ')[0].split('.')[1]),
                                           int(period_mask['start'][k].split(' ')[0].split('.')[0]),
                                           int(period_mask['start'][k].split(' ')[1].split(':')[0]),
                                           int(period_mask['start'][k].split(' ')[1].split(':')[1]),
                                           int(period_mask['start'][k].split(' ')[1].split(':')[2])))
            period_mask['start_date'] = start_date
            period_mask['end_date'] = end_date

            # Непонятно как делать отдельно шумы и нейтроны
            for t in range(len(mask_without_periods.index)):
                merge_utc.loc[merge_utc[(merge_utc['DATE'].isin([mask_without_periods['start_date'][t]])) & (
                    merge_utc['time'].isin([mask_without_periods['start_time'][t]]))].index, 'Nn%s' % i] = 0
            for t in range(len(mask_without_periods.index)):
                merge_utc.loc[merge_utc[(merge_utc['DATE'].isin([mask_without_periods['start_date'][t]])) & (
                    merge_utc['time'].isin([mask_without_periods['start_time'][t]]))].index, 'N_noise%s' % i] = 0

            for a in range(len(period_mask.index)):
                merge_utc['Nn%s' % i] = merge_utc['Nn%s' % i].where(
                    (merge_utc['new_format_time'] < period_mask['start_date'][a]) | (
                            merge_utc['new_format_time'] > period_mask['end_date'][a]), 0)
            for a in range(len(period_mask.index)):
                merge_utc['N_noise%s' % i] = merge_utc['N_noise%s' % i].where(
                    (merge_utc['new_format_time'] < period_mask['start_date'][a]) | (
                            merge_utc['new_format_time'] > period_mask['end_date'][a]), 0)

    except FileNotFoundError:
        mask_checker = True
        print('Mask-file dont exist')

    if len(mean_pressure_uragan) >= len(merge_utc.index):
        mean_pressure_uragan = mean_pressure_uragan[0:len(merge_utc.index)]

    def make_corr_data(type_of_impulse, pressure, type_of_impulse_sign, pathpic, styear, stday, stmonth, endday,
                       endmonth):
        fig, axs = plt.subplots(figsize=(18, 18), nrows=4, sharex=True)
        N_bar_koef = []
        N_0 = []
        B_koef = []
        for i in range(1, 5):
            ax = axs[i - 1]
            ax.set_title('Детектор %s' % i, fontsize=18, loc='left')
            if i == 5:
                ax.set_xlabel('Давление', fontsize=20)
            ax.set_ylabel(type_of_impulse_sign + r'$, (300с)^{-1}$', fontsize=15)
            #     ax.set_ylim([0,500])
            #     ax.set_xlim([0,df.index.max()])
            if len(merge_utc[merge_utc[type_of_impulse + '%s' % i] == 0]) != 0:
                test = []
                for j in merge_utc[merge_utc[type_of_impulse + '%s' % i] != 0].index.tolist():
                    if mean_pressure_uragan[j] != 0:
                        test.append(mean_pressure_uragan[j])
                    else:
                        merge_utc.loc[j, type_of_impulse + '%s' % i] = 0
                ax.scatter([x - pressure for x in test],
                           merge_utc[merge_utc[type_of_impulse + '%s' % i] != 0][type_of_impulse + '%s' % i],
                           label=type_of_impulse + '%s' % i, s=5)
                param, param_cov = curve_fit(linear, [x - pressure for x in test],
                                             merge_utc[merge_utc[type_of_impulse + '%s' % i] != 0][
                                                 type_of_impulse + '%s' % i])
                fit_line = [y * param[0] + param[1] for y in [x - pressure for x in test]]
                ax.scatter([x - pressure for x in test], fit_line, s=6)
                N_bar_koef.append(param[0] / param[1] * 100)
                N_0.append(param[1])
                B_koef.append(param[0])
            else:
                ax.scatter([x - pressure for x in mean_pressure_uragan], merge_utc[type_of_impulse + '%s' % i],
                           label=type_of_impulse + '%s' % i, s=5)
                param, param_cov = curve_fit(linear, [x - pressure for x in mean_pressure_uragan],
                                             merge_utc[type_of_impulse + '%s' % i])
                fit_line = [y * param[0] + param[1] for y in [x - pressure for x in mean_pressure_uragan]]
                ax.scatter([x - pressure for x in mean_pressure_uragan], fit_line, s=6)
                N_bar_koef.append(param[0] / param[1] * 100)
                N_0.append(param[1])
                B_koef.append(param[0])

        plt.savefig(
            '{}\\{}\\{}(P){:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, type_of_impulse, stday, stmonth,
                                                              endday, endmonth),
            bbox_inches='tight')

        press_for_fixed_N = []
        for i in range(4):
            intermediate_corr = []
            for x in mean_pressure_uragan:
                if x != 0:
                    intermediate_corr.append((x - pressure) * N_bar_koef[i] / 100 * N_0[i])
                else:
                    intermediate_corr.append(0)
            press_for_fixed_N.append(intermediate_corr)

        return press_for_fixed_N, N_bar_koef, B_koef

    if proccessing_pressure == 'mean_value':
        press_for_fixed_N, N_bar_koef, B_koef = make_corr_data('Nn', np.mean(mean_pressure_uragan), 'Скорость счета',
                                                               pathpic,
                                                               styear,
                                                               stday,
                                                               stmonth, endday,
                                                               endmonth)
        press_for_fixed_N_noise, N_noise_bar_koef, B_noise_koef = make_corr_data('N_noise',
                                                                                 np.mean(mean_pressure_uragan),
                                                                                 'Скорость счета шумов',
                                                                                 pathpic,
                                                                                 styear, stday,
                                                                                 stmonth, endday,
                                                                                 endmonth)
    else:
        press_for_fixed_N, N_bar_koef, B_koef = make_corr_data('Nn', 993, 'Скорость счета', pathpic, styear,
                                                               stday,
                                                               stmonth, endday,
                                                               endmonth)
        press_for_fixed_N_noise, N_noise_bar_koef, B_noise_koef = make_corr_data('N_noise', 993, 'Скорость счета шумов',
                                                                                 pathpic,
                                                                                 styear, stday,
                                                                                 stmonth, endday,
                                                                                 endmonth)

    Nn_to_P_path = "{}\\{}\\Nn(P){:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)
    Nnoise_to_P_path = "{}\\{}\\N_noise(P){:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth, endday,
                                                                              endmonth)
    fig, axs = plt.subplots(figsize=(18, 18), nrows=4, sharex=True)
    corr_N = []
    for i in range(1, 5):
        merge_utc['corr_N%s' % i] = merge_utc['Nn%s' % i] - press_for_fixed_N[i - 1]
        merge_utc['corr_N%s' % i].where(merge_utc['corr_N%s' % i] > 10, 0, inplace=True)
        merge_utc['corr_N%s' % i].where(merge_utc['corr_N%s' % i] != merge_utc['Nn%s' % i], 0, inplace=True)
        corr_N.append(merge_utc['corr_N%s' % i])
        ax = axs[i - 1]
        if len(merge_utc[merge_utc['Nn%s' % i] == 0]) != 0:
            y_min = np.min(merge_utc[merge_utc['Nn%s' % i] > 5]['Nn%s' % i]) - 30
            y_max = np.max(merge_utc[merge_utc['Nn%s' % i] > 5]['Nn%s' % i]) + 30
            ax.set_ylim([y_min, y_max])
        ax.set_title('Детектор %s' % i, fontsize=18, loc='left')
        ax.grid()
        ax0 = ax.twinx()
        ax0.set_ylim([970, 1020])
        ax0.set_ylabel('Давление, мбар', fontsize=18)
        if i == 4:
            ax.set_xlabel('Дата', fontsize=20)
        ax.set_ylabel('Cкорость счета, (300с)⁻¹', fontsize=16)
        ax.set_xlim([0, merge_utc.index.max()])
        ax.minorticks_on()
        ax.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
        ax.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
        ax.plot(merge_utc.index, merge_utc['Nn%s' % i], label='N%s' % i, linewidth=1, color='black')
        ax.plot(merge_utc.index, merge_utc['corr_N%s' % i], label='N%s' % i, linewidth=1,
                color='red')

        ax0.plot(range(0, len(mean_pressure_uragan)), mean_pressure_uragan, linewidth=2, color='blue')
        ax.set_xticks(list(range(0, merge_utc.index.max(), 288 * 4)))
        ax.set_xticklabels(merge_utc['DATE'].dt.date.unique().tolist()[::4])
    plt.savefig('{}\\{}\\Nn300c{:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    Nn_path = "{}\\{}\\Nn300c{:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)

    fig, axs = plt.subplots(figsize=(18, 18), nrows=4, sharex=True)
    corr_N_noise = []
    for i in range(1, 5):
        merge_utc['corr_N_noise%s' % i] = merge_utc['N_noise%s' % i] - press_for_fixed_N_noise[i - 1]
        merge_utc['corr_N_noise%s' % i].where(merge_utc['corr_N_noise%s' % i] > 10, 0, inplace=True)
        merge_utc['corr_N%s' % i].where(merge_utc['corr_N_noise%s' % i] != merge_utc['N_noise%s' % i], 0, inplace=True)
        corr_N_noise.append(merge_utc['corr_N_noise%s' % i])
        ax = axs[i - 1]
        ax.set_title('Детектор %s' % i, fontsize=18, loc='left')
        if len(merge_utc[merge_utc['N_noise%s' % i] == 0]) != 0:
            y_min = np.min(merge_utc[merge_utc['N_noise%s' % i] > 5]['N_noise%s' % i]) - 30
            y_max = np.max(merge_utc[merge_utc['N_noise%s' % i] > 5]['N_noise%s' % i]) + 30
            ax.set_ylim([y_min, y_max])
        ax.grid()
        ax0 = ax.twinx()
        ax0.set_ylim([970, 1020])
        ax0.set_ylabel('Давление, мбар', fontsize=18)
        if i == 4:
            ax.set_xlabel('Дата', fontsize=20)
        ax.set_ylabel('Cкорость счета шумов, (300с)⁻¹', fontsize=14)

        ax.set_xlim([0, merge_utc.index.max()])
        ax.minorticks_on()
        ax.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
        ax.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
        ax.plot(merge_utc.index, merge_utc['N_noise%s' % i], label='N%s' % i, linewidth=1, color='black')
        ax.plot(merge_utc.index, merge_utc['corr_N_noise%s' % i], label='N%s' % i,
                linewidth=1,
                color='red')
        ax0.plot(range(0, len(mean_pressure_uragan)), mean_pressure_uragan, linewidth=2, color='blue')
        ax.set_xticks(list(range(0, merge_utc.index.max(), 288 * 4)))
        ax.set_xticklabels(merge_utc['DATE'].dt.date.unique().tolist()[::4])
    plt.savefig(
        '{}\\{}\\Nnoise300c{:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
        bbox_inches='tight')
    Nnoise_path = "{}\\{}\\Nnoise300c{:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth, endday,
                                                                         endmonth)

    frame = merge_utc[
        (merge_utc['Nn1'] != 0) & (merge_utc['Nn2'] != 0) & (merge_utc['Nn3'] != 0) & (merge_utc['Nn4'] != 0)]

    counter = 0
    b = date(endyear, endmonth, endday)
    while counter != 288:
        b = b - timedelta(1)
        counter = len(frame[frame['DATE'].dt.date == b])

    distribution_cols = []
    for i in range(4):
        distribution_cols.append("det%s" % (i + 1))

    # Прописать рандомную функцию выбора даты
    # Заменить название рисунков пути к распределениям на даты из рандомной функции
    try:
        R_distribution = pd.read_csv(
            '{}\\sp\\4R{:02}-{:02}.{:02}'.format(file1cl,
                                                 b.month,
                                                 b.day,
                                                 b.year - 2000),
            sep=' ', header=None, skipinitialspace=True, index_col=0)
        R_distribution = R_distribution.dropna(axis=1, how='all')
        R_distribution.columns = distribution_cols

        plt.figure(figsize=(18, 10))
        plt.xlabel('R', fontsize=20)
        plt.ylabel('Число событий', fontsize=20)
        plt.xlim([0, 100])
        plt.ylim([0, 12000])
        plt.grid()
        plt.minorticks_on()
        box_1 = {'facecolor': 'white',  # цвет области
                 'edgecolor': 'red',  # цвет крайней линии
                 'boxstyle': 'round'}
        plt.text(43, 12000, "за {:02}.{:02}.{:02}".format(b.day, b.month, b.year - 2000), bbox=box_1, fontsize=20)
        plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
        plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
        plt.grid(which='minor',
                 color='k',
                 linestyle=':')
        for i in range(1, 5):
            plt.plot(R_distribution.index, R_distribution['det%s' % i], label='Детектор %s' % i, linewidth=4.5)
        plt.legend(loc="upper right")
        plt.savefig(
            '{}\\{}\\R_distribution{:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, stday, stmonth, endday,
                                                                       endmonth),
            bbox_inches='tight')
        R_distribution_path = "{}\\{}\\R_distribution{:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday,
                                                                                         stmonth,
                                                                                         endday,
                                                                                         endmonth)
    except:
        print("такого файла нит 4R{:02}-{:02}.{:02}".format(
            b.month,
            b.day,
            b.year - 2000))
    try:
        Front_time = pd.read_csv(
            '{}\\sp\\4Tf{:02}-{:02}.{:02}'.format(file1cl,
                                                  b.month,
                                                  b.day,
                                                  b.year - 2000),
            sep=' ', header=None, skipinitialspace=True, index_col=0)
        Front_time = Front_time.dropna(axis=1, how='all')
        Front_time.columns = distribution_cols

        plt.figure(figsize=(18, 10))
        plt.xlabel('Время нарастаний фронта', fontsize=20)
        plt.ylabel('Число событий', fontsize=20)
        plt.yscale('log')
        plt.xlim([0, 200])
        plt.ylim([1, 10000])
        plt.grid()
        plt.minorticks_on()
        plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
        plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
        box_1 = {'facecolor': 'white',  # цвет области
                 'edgecolor': 'red',  # цвет крайней линии
                 'boxstyle': 'round'}
        plt.text(87, 10000, "за {:02}.{:02}.{:02}".format(b.day, b.month, b.year - 2000), bbox=box_1, fontsize=20)
        plt.grid(which='minor',
                 color='k',
                 linestyle=':')
        for i in range(1, 5):
            plt.plot(Front_time.index, Front_time['det%s' % i], label='Детектор %s' % i, linewidth=4.5)
        plt.legend(loc="upper right")
        plt.savefig(
            '{}\\{}\\Front_time{:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
            bbox_inches='tight')
        Front_time_path = "{}\\{}\\Front_time{:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth,
                                                                                 endday,
                                                                                 endmonth)
    except:
        print("такого файла нит 4Tf{:02}-{:02}.{:02}".format(
            b.month,
            b.day,
            b.year - 2000))

    try:
        distribution_cols = []
        for i in range(4):
            distribution_cols.append("det%s" % (i + 1))
        for i in range(4):
            distribution_cols.append("noise%s" % (i + 1))

        N_amp = pd.read_csv(
            '{}\\sp\\4sp{:02}-{:02}.{:02}'.format(file1cl,
                                                  b.month,
                                                  b.day,
                                                  b.year - 2000),
            sep=' ', header=None, skipinitialspace=True, index_col=0)
        N_amp = N_amp.dropna(axis=1, how='all')
        N_amp.columns = distribution_cols

        plt.figure(figsize=(18, 10))
        plt.xlabel('Амплитуда', fontsize=20)
        plt.ylabel('Число событий', fontsize=20)
        plt.yscale('log')
        plt.xlim([0, 250])
        plt.ylim([1, 10000])
        plt.grid()
        plt.minorticks_on()
        plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
        plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
        plt.grid(which='minor',
                 color='k',
                 linestyle=':')
        Nline = []
        Noiseline = []
        for i in range(1, 5):
            n_line, = plt.plot(N_amp.index, N_amp['det%s' % i], label='Детектор %s' % i, linewidth=4.5)
            Nline.append(n_line)
        for i in range(1, 5):
            noise_line, = plt.plot(N_amp.index, N_amp['noise%s' % i], label='Детектор %s' % i, linewidth=4.5)
            Noiseline.append(noise_line)
        box_1 = {'facecolor': 'white',  # цвет области
                 'edgecolor': 'red',  # цвет крайней линии
                 'boxstyle': 'round'}
        plt.text(110, 10000, "за {:02}.{:02}.{:02}".format(b.day, b.month, b.year - 2000), bbox=box_1, fontsize=20)
        first_legend = plt.legend(handles=Nline, loc='upper center', bbox_to_anchor=(0.75, 1), title='Нейтроны')
        plt.gca().add_artist(first_legend)
        plt.legend(handles=Noiseline, loc='upper right', title='Шумы')
        plt.savefig('{}\\{}\\N_amp{:02}-{:02}-{:02}-{:02}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                    bbox_inches='tight')
        N_amp_path = "{}\\{}\\N_amp{:02}-{:02}-{:02}-{:02}.png".format(pathpic, styear, stday, stmonth, endday,
                                                                       endmonth)

    except:
        print("такого файла нит 4sp{:02}-{:02}.{:02}".format(
            b.month,
            b.day,
            b.year - 2000))

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    styles = doc.styles
    style = styles.add_style('PItalic', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['PItalic']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.italic = True
    font.bold = True

    headstyle = styles.add_style('Headstyle', WD_STYLE_TYPE.PARAGRAPH)
    headstyle = doc.styles['Headstyle']
    font = headstyle.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True

    headgraf = styles.add_style('Headgraf', WD_STYLE_TYPE.PARAGRAPH)
    headgraf = doc.styles['Headgraf']
    font = headgraf.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.bold = True
    font.italic = True

    head = doc.add_paragraph(
        'Справка о работе установки «Нейтрон» в период с {:02}.{:02}.{} по {:02}.{:02}.{} '.format(stday, stmonth,
                                                                                                   styear,
                                                                                                   endday, endmonth,
                                                                                                   endyear),
        style='Headstyle')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    desc = doc.add_paragraph('Таблица 1: Календарная эффективность.', style='PItalic')

    beg = doc.add_table(5, 4, doc.styles['Table Grid'])
    # beg.alignment = WD_TABLE_ALIGNMENT.CENTER
    beg.cell(0, 0).text = '№ детектора'
    beg.cell(0, 1).text = 'Экспозиции, ч.'
    beg.cell(0, 2).text = 'Календарное время, ч.'
    beg.cell(0, 3).text = 'Экспозиция, %'
    beg.cell(1, 0).text = '1'
    beg.cell(1, 1).text = str(round(worktime['WORKTIME1'].sum(), 2))
    beg.cell(1, 2).text = str(24 * (len(worktime)))
    beg.cell(1, 3).text = str(round(worktime['WORKTIME1'].sum() / (24 * (len(worktime))) * 100, 3)) + '%'
    beg.cell(2, 0).text = '2'
    beg.cell(2, 1).text = str(round(worktime['WORKTIME2'].sum(), 2))
    beg.cell(2, 2).text = str(24 * (len(worktime)))
    beg.cell(2, 3).text = str(round(worktime['WORKTIME2'].sum() / (24 * (len(worktime))) * 100, 3)) + '%'
    beg.cell(3, 0).text = '3'
    beg.cell(3, 1).text = str(round(worktime['WORKTIME3'].sum(), 2))
    beg.cell(3, 2).text = str(24 * (len(worktime)))
    beg.cell(3, 3).text = str(round(worktime['WORKTIME3'].sum() / (24 * (len(worktime))) * 100, 3)) + '%'
    beg.cell(4, 0).text = '4'
    beg.cell(4, 1).text = str(round(worktime['WORKTIME4'].sum(), 2))
    beg.cell(4, 2).text = str(24 * (len(worktime)))
    beg.cell(4, 3).text = str(round(worktime['WORKTIME4'].sum() / (24 * (len(worktime))) * 100, 3)) + '%'

    for row in range(1):
        for col in range(5):
            # получаем ячейку таблицы
            cell = beg.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, 5):
        for col in range(1):
            # получаем ячейку таблицы
            cell = beg.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(5):
        for col in range(4):
            cell = beg.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    desc = doc.add_paragraph('Таблица 2: Сводная таблица неисправностей установки.', style='PItalic')

    # Переработать близко с предыдущей таблицей
    err = doc.add_table(len(failstr_begin) + 2, 5, doc.styles['Table Grid'])
    # err.alignment = WD_TABLE_ALIGNMENT.CENTER
    err.cell(0, 0).text = '№'
    err.cell(0, 0).merge(err.cell(1, 0))
    err.cell(0, 1).text = 'Время простоя'
    err.cell(1, 1).text = 'c'
    err.cell(1, 2).text = 'по'
    err.cell(0, 1).merge(err.cell(0, 2))
    err.cell(0, 3).text = 'Кол-во потерянных минут (период)'
    err.cell(0, 3).merge(err.cell(1, 3))
    err.cell(0, 4).text = 'Примечание'
    err.cell(0, 4).merge(err.cell(1, 4))

    for i in range(2, len(failstr_begin) + 2):
        err.cell(i, 0).text = str(timestop['claster'][i - 2])
        err.cell(i, 1).text = str(failstr_begin[i - 2])
        err.cell(i, 2).text = str(failstr_end[i - 2])
        err.cell(i, 3).text = str(lost_minutes[i - 2])
        err.cell(i, 4).text = ' '

    for row in range(2):
        for col in range(5):
            # получаем ячейку таблицы
            cell = err.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, len(failstr_begin) + 2):
        for col in range(1):
            # получаем ячейку таблицы
            cell = err.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(len(failstr_begin) + 2):
        for col in range(5):
            cell = err.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    desc = doc.add_paragraph(
        r'Таблица 3: Средняя скорость счета [(300с)⁻¹]',  # возвести в минус первую степень #сделано!
        style='PItalic')

    neut = doc.add_table(3, 5, doc.styles['Table Grid'])
    # neut.alignment = WD_TABLE_ALIGNMENT.CENTER
    neut.cell(0, 0).text = '№ детектора'
    neut.cell(0, 1).text = '1'
    neut.cell(0, 2).text = '2'
    neut.cell(0, 3).text = '3'
    neut.cell(0, 4).text = '4'
    neut.cell(1, 0).text = 'Скорость счета нейтронных импульсов, (300с)⁻¹'
    neut.cell(1, 1).text = str(round(merge_utc[merge_utc['Nn1'] != 0]['Nn1'].mean(), 2))
    neut.cell(1, 2).text = str(round(merge_utc[merge_utc['Nn2'] != 0]['Nn2'].mean(), 2))
    neut.cell(1, 3).text = str(round(merge_utc[merge_utc['Nn3'] != 0]['Nn3'].mean(), 2))
    neut.cell(1, 4).text = str(round(merge_utc[merge_utc['Nn4'] != 0]['Nn4'].mean(), 2))
    neut.cell(2, 0).text = 'Скорость счета шумовых импульсов, (300с)⁻¹'
    neut.cell(2, 1).text = str(round(merge_utc[merge_utc['N_noise1'] != 0]['N_noise1'].mean(), 2))
    neut.cell(2, 2).text = str(round(merge_utc[merge_utc['N_noise2'] != 0]['N_noise2'].mean(), 2))
    neut.cell(2, 3).text = str(round(merge_utc[merge_utc['N_noise3'] != 0]['N_noise3'].mean(), 2))
    neut.cell(2, 4).text = str(round(merge_utc[merge_utc['N_noise4'] != 0]['N_noise4'].mean(), 2))

    for row in range(1):
        for col in range(5):
            # получаем ячейку таблицы
            cell = neut.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, 3):
        for col in range(1):
            # получаем ячейку таблицы
            cell = neut.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(3):
        for col in range(1, 5):
            cell = neut.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in neut.columns[0].cells:
        cell.width = Inches(3)
    for cell in neut.columns[1].cells:
        cell.width = Inches(1.075)
    for cell in neut.columns[2].cells:
        cell.width = Inches(1.075)
    for cell in neut.columns[3].cells:
        cell.width = Inches(1.075)
    for cell in neut.columns[4].cells:
        cell.width = Inches(1.075)

    space = doc.add_paragraph()

    desc = doc.add_paragraph(
        'Таблица 4: Барометрические коэффициенты β ',
        style='PItalic')

    baro = doc.add_table(3, 5, doc.styles['Table Grid'])
    # baro.alignment = WD_TABLE_ALIGNMENT.CENTER
    baro.cell(0, 0).text = '№ детектора'
    baro.cell(0, 1).text = '1'
    baro.cell(0, 2).text = '2'
    baro.cell(0, 3).text = '3'
    baro.cell(0, 4).text = '4'
    baro.cell(1, 0).text = 'β нейтронных импульсов, %/мбар'
    for i in range(1, 5):
        baro.cell(1, i).text = str(round(N_bar_koef[i - 1], 2))
    baro.cell(2, 0).text = 'β шумовых импульсов, %/мбар'
    for i in range(1, 5):
        baro.cell(2, i).text = str(round(N_noise_bar_koef[i - 1], 2))

    for row in range(1):
        for col in range(5):
            # получаем ячейку таблицы
            cell = baro.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in range(1, 3):
        for col in range(1):
            # получаем ячейку таблицы
            cell = baro.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in range(3):
        for col in range(1, 5):
            cell = baro.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in baro.columns[0].cells:
        cell.width = Inches(3)
    for cell in baro.columns[1].cells:
        cell.width = Inches(1.075)
    for cell in baro.columns[2].cells:
        cell.width = Inches(1.075)
    for cell in baro.columns[3].cells:
        cell.width = Inches(1.075)
    for cell in baro.columns[4].cells:
        cell.width = Inches(1.075)

    space = doc.add_paragraph()

    desc = doc.add_paragraph(
        'Таблица 5: Коеффициент B',
        style='PItalic')

    B = doc.add_table(3, 5, doc.styles['Table Grid'])
    # B.alignment = WD_TABLE_ALIGNMENT.CENTER
    B.cell(0, 0).text = '№ детектора'
    B.cell(0, 1).text = '1'
    B.cell(0, 2).text = '2'
    B.cell(0, 3).text = '3'
    B.cell(0, 4).text = '4'
    B.cell(1, 0).text = 'B нейтронных импульсов'
    print(B_koef)
    print(B_noise_koef)
    for i in range(1, 5):
        B.cell(1, i).text = str(round(B_koef[i - 1], 2))
    B.cell(2, 0).text = 'B шумовых импульсов'
    for i in range(1, 5):
        B.cell(2, i).text = str(round(B_noise_koef[i - 1], 2))

    for row in range(1):
        for col in range(5):
            # получаем ячейку таблицы
            cell = B.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in range(1, 3):
        for col in range(1):
            # получаем ячейку таблицы
            cell = B.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in range(3):
        for col in range(1, 5):
            cell = B.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in B.columns[0].cells:
        cell.width = Inches(3)
    for cell in B.columns[1].cells:
        cell.width = Inches(1.075)
    for cell in B.columns[2].cells:
        cell.width = Inches(1.075)
    for cell in B.columns[3].cells:
        cell.width = Inches(1.075)
    for cell in B.columns[4].cells:
        cell.width = Inches(1.075)

    space = doc.add_paragraph()
    if (mask_checker):
        desc = doc.add_paragraph('Выборка данных с помощью маски не была произведена',
                                 style='Headstyle')

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Зависимости скорости счета нейтронных импульсов и давления от времени.', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    doc.add_picture(Nn_path, width=Inches(7.5), height=Inches(7.9))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Зависимости скорости счета шумовых импульсов и давления от времени', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    doc.add_picture(Nnoise_path, width=Inches(7.5), height=Inches(7.9))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Зависимости скорости счета нейтронных импульсов от давления', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    doc.add_picture(Nn_to_P_path, width=Inches(7.5), height=Inches(7.9))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Зависимости скорости счета шумовых импульсов от давления', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    space = doc.add_paragraph()

    doc.add_picture(Nnoise_to_P_path, width=Inches(7.5), height=Inches(7.9))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph(
        'Временные распределения сигналов',
        style='Headgraf')

    space = doc.add_paragraph()

    try:
        doc.add_picture(N_amp_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        desc = doc.add_paragraph(
            'Амплитудные распределения сигналов',
            style='Headgraf')

        space = doc.add_paragraph()

        doc.add_picture(Front_time_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = doc.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)

        desc = doc.add_paragraph(
            'Распределения сигналов по параметру R (Af/Amax)',
            style='Headgraf')

        space = doc.add_paragraph()

        doc.add_picture(R_distribution_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        print("Не было данных о графике распределения")

    add_page_number(doc.sections[0].footer.paragraphs[0])

    doc.save(f'{path}\{stday:02}.{stmonth:02}.{styear}-{endday:02}.{endmonth:02}.{endyear}.docx')
